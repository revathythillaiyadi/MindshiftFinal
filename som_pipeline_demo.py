#!/usr/bin/env python3
"""
MindShift SOM Data Loading and Embedding Pipeline - DEMO
Complete working implementation using LangChain

This script demonstrates the complete data loading and embedding pipeline
that mirrors the n8n 'Load Data Flow' functionality.
"""

import os
import logging
from pathlib import Path
from typing import List, Dict, Any
from dotenv import load_dotenv

# LangChain imports
from langchain_text_splitters import RecursiveCharacterTextSplitter
from langchain_openai import OpenAIEmbeddings
from langchain_community.vectorstores import Chroma
from langchain_core.documents import Document

# Configure logging
logging.basicConfig(
    level=logging.INFO,
    format='%(asctime)s - %(levelname)s - %(message)s'
)
logger = logging.getLogger(__name__)

def load_docx_documents(docs_directory: str) -> List[Document]:
    """Load all .docx files from the documents directory"""
    logger.info(f"📚 Loading documents from: {docs_directory}")
    
    docs_path = Path(docs_directory)
    if not docs_path.exists():
        raise FileNotFoundError(f"Documents directory not found: {docs_directory}")
    
    documents = []
    
    for docx_file in docs_path.glob("*.docx"):
        try:
            from docx import Document as DocxDocument
            
            doc = DocxDocument(docx_file)
            content_parts = []
            
            for paragraph in doc.paragraphs:
                if paragraph.text.strip():
                    content_parts.append(paragraph.text.strip())
            
            content = "\n".join(content_parts)
            
            if content.strip():
                doc_obj = Document(
                    page_content=content,
                    metadata={"source": str(docx_file), "filename": docx_file.name}
                )
                documents.append(doc_obj)
                logger.info(f"✅ {docx_file.name}: {len(content)} characters")
            
        except Exception as e:
            logger.error(f"❌ Error loading {docx_file.name}: {e}")
    
    logger.info(f"📊 Total documents loaded: {len(documents)}")
    return documents

def split_documents(documents: List[Document], chunk_size: int = 1000, chunk_overlap: int = 200) -> List[Document]:
    """Split documents into chunks using RecursiveCharacterTextSplitter"""
    logger.info("✂️ Splitting documents into chunks...")
    
    text_splitter = RecursiveCharacterTextSplitter(
        chunk_size=chunk_size,
        chunk_overlap=chunk_overlap,
        length_function=len,
        separators=["\n\n", "\n", " ", ""]
    )
    
    split_docs = text_splitter.split_documents(documents)
    
    # Calculate statistics
    chunk_sizes = [len(chunk.page_content) for chunk in split_docs]
    avg_size = sum(chunk_sizes) / len(chunk_sizes)
    
    logger.info(f"📊 Split into {len(split_docs)} chunks")
    logger.info(f"📊 Average chunk size: {avg_size:.0f} characters")
    logger.info(f"📊 Chunk size range: {min(chunk_sizes)} - {max(chunk_sizes)} characters")
    
    return split_docs

def create_embeddings_and_vectorstore(documents: List[Document], 
                                    collection_name: str = "som_mindshift",
                                    persist_directory: str = "./chroma_db") -> Chroma:
    """Create OpenAI embeddings and ChromaDB vector store"""
    logger.info("🔮 Creating embeddings and vector store...")
    
    # Initialize OpenAI Embeddings
    embeddings = OpenAIEmbeddings(
        model="text-embedding-3-small",
        openai_api_key=os.getenv("OPENAI_API_KEY")
    )
    
    logger.info("✅ OpenAI Embeddings initialized")
    
    # Create persistent directory
    persist_path = Path(persist_directory)
    persist_path.mkdir(exist_ok=True)
    
    # Create ChromaDB vector store
    vectorstore = Chroma.from_documents(
        documents=documents,
        embedding=embeddings,
        collection_name=collection_name,
        persist_directory=str(persist_path)
    )
    
    # Persist the vector store
    vectorstore.persist()
    
    logger.info(f"✅ Vector store created: {collection_name}")
    logger.info(f"✅ Persisted to: {persist_directory}")
    
    return vectorstore

def get_collection_stats(vectorstore: Chroma) -> Dict[str, Any]:
    """Get statistics about the created collection"""
    try:
        collection = vectorstore._collection
        count = collection.count()
        
        stats = {
            "collection_name": vectorstore._collection.name,
            "total_documents": count,
            "embedding_model": "text-embedding-3-small",
            "persist_directory": str(vectorstore._persist_directory)
        }
        
        return stats
    except Exception as e:
        logger.error(f"Error getting stats: {e}")
        return {"error": str(e)}

def main():
    """Main function demonstrating the complete pipeline"""
    print("🧠 MindShift SOM Data Loading and Embedding Pipeline")
    print("=" * 60)
    print("Mirrors the n8n 'Load Data Flow' using LangChain")
    print("=" * 60)
    
    # Load environment variables
    load_dotenv()
    
    # Check for OpenAI API key
    if not os.getenv("OPENAI_API_KEY"):
        print("❌ OPENAI_API_KEY environment variable not set")
        print("Please set your OpenAI API key in a .env file")
        print("You can get an API key from: https://platform.openai.com/api-keys")
        return
    
    try:
        # Step 1: Load documents
        print("\n🚀 Step 1: Loading Documents")
        print("-" * 30)
        documents = load_docx_documents("./som_documents")
        
        # Step 2: Split documents
        print("\n✂️ Step 2: Splitting Documents")
        print("-" * 30)
        split_docs = split_documents(documents, chunk_size=1000, chunk_overlap=200)
        
        # Step 3: Create embeddings and vector store
        print("\n🔮 Step 3: Creating Embeddings and Vector Store")
        print("-" * 30)
        vectorstore = create_embeddings_and_vectorstore(
            split_docs,
            collection_name="som_mindshift",
            persist_directory="./chroma_db"
        )
        
        # Step 4: Get statistics
        print("\n📊 Step 4: Collection Statistics")
        print("-" * 30)
        stats = get_collection_stats(vectorstore)
        
        if "error" not in stats:
            print(f"✅ Collection: {stats['collection_name']}")
            print(f"✅ Total documents: {stats['total_documents']}")
            print(f"✅ Embedding model: {stats['embedding_model']}")
            print(f"✅ Persist directory: {stats['persist_directory']}")
        
        # Success message
        print("\n" + "=" * 60)
        print("🎉 PIPELINE COMPLETED SUCCESSFULLY!")
        print("=" * 60)
        print(f"📊 Documents processed: {len(documents)}")
        print(f"📊 Chunks created: {len(split_docs)}")
        print(f"📊 Vector store: som_mindshift")
        print(f"📊 Embeddings: text-embedding-3-small")
        
        print("\n🎯 Next Steps:")
        print("1. Use the vector store for similarity search")
        print("2. Implement RAG query functionality")
        print("3. Create a chatbot interface")
        
        print("\n💡 Example usage:")
        print("```python")
        print("from langchain_community.vectorstores import Chroma")
        print("from langchain_openai import OpenAIEmbeddings")
        print("")
        print("# Load the vector store")
        print("embeddings = OpenAIEmbeddings(model='text-embedding-3-small')")
        print("vectorstore = Chroma(")
        print("    collection_name='som_mindshift',")
        print("    embedding_function=embeddings,")
        print("    persist_directory='./chroma_db'")
        print(")")
        print("")
        print("# Perform similarity search")
        print("results = vectorstore.similarity_search('limiting belief', k=5)")
        print("```")
        
    except Exception as e:
        logger.error(f"❌ Pipeline failed: {e}")
        print(f"\n❌ Error: {e}")

if __name__ == "__main__":
    main()

