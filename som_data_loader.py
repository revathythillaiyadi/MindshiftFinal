#!/usr/bin/env python3
"""
MindShift SOM Data Loading and Embedding Pipeline
Mirrors the n8n 'Load Data Flow' using LangChain

This script performs:
1. File Loading: Load and read all .docx files (01 through 08)
2. Text Splitting: Use RecursiveCharacterTextSplitter for optimal chunking
3. Embedding: Initialize OpenAI Embeddings model
4. Vector Store: Create persistent ChromaDB collection

Based on documents from: https://drive.google.com/drive/folders/1k9rsWGuIjYBTqmXpE0FKP8LJcnWuXHgp
"""

import os
import logging
from pathlib import Path
from typing import List, Dict, Any
from dotenv import load_dotenv

# LangChain imports
from langchain_community.document_loaders import DirectoryLoader
from langchain_text_splitters import RecursiveCharacterTextSplitter
from langchain_openai import OpenAIEmbeddings
from langchain_community.vectorstores import Chroma
from langchain_core.documents import Document

# Configure logging
logging.basicConfig(
    level=logging.INFO,
    format='%(asctime)s - %(levelname)s - %(message)s'
)
logger = logging.getLogger(__name__)

class SOMDataLoader:
    """
    MindShift SOM Data Loading and Embedding Pipeline
    Mirrors the n8n 'Load Data Flow' functionality
    """
    
    def __init__(self, 
                 docs_directory: str = "./som_documents",
                 collection_name: str = "som_mindshift",
                 persist_directory: str = "./chroma_db",
                 chunk_size: int = 1000,
                 chunk_overlap: int = 200):
        """
        Initialize the SOM Data Loader
        
        Args:
            docs_directory: Path to directory containing .docx files
            collection_name: Name for ChromaDB collection
            persist_directory: Directory to persist ChromaDB data
            chunk_size: Size of text chunks for splitting
            chunk_overlap: Overlap between chunks
        """
        self.docs_directory = Path(docs_directory)
        self.collection_name = collection_name
        self.persist_directory = Path(persist_directory)
        self.chunk_size = chunk_size
        self.chunk_overlap = chunk_overlap
        
        # Initialize components
        self.text_splitter = None
        self.embeddings = None
        self.vectorstore = None
        
        # Load environment variables
        load_dotenv()
        
        # Validate OpenAI API key
        if not os.getenv("OPENAI_API_KEY"):
            raise ValueError("OPENAI_API_KEY environment variable not set")
    
    def setup_text_splitter(self) -> None:
        """Initialize the RecursiveCharacterTextSplitter"""
        logger.info("Setting up RecursiveCharacterTextSplitter...")
        
        self.text_splitter = RecursiveCharacterTextSplitter(
            chunk_size=self.chunk_size,
            chunk_overlap=self.chunk_overlap,
            length_function=len,
            separators=[
                "\n\n",  # Split on double newlines (paragraphs)
                "\n",   # Split on single newlines
                " ",    # Split on spaces
                "",     # Split on characters
            ]
        )
        
        logger.info(f"✅ Text splitter configured: chunk_size={self.chunk_size}, overlap={self.chunk_overlap}")
    
    def setup_embeddings(self) -> None:
        """Initialize OpenAI Embeddings model"""
        logger.info("Setting up OpenAI Embeddings...")
        
        self.embeddings = OpenAIEmbeddings(
            model="text-embedding-3-small",
            openai_api_key=os.getenv("OPENAI_API_KEY")
        )
        
        logger.info("✅ OpenAI Embeddings initialized")
    
    def load_documents(self) -> List[Document]:
        """
        Load all .docx files from the documents directory
        
        Returns:
            List of Document objects
        """
        logger.info(f"Loading documents from: {self.docs_directory}")
        
        if not self.docs_directory.exists():
            raise FileNotFoundError(f"Documents directory not found: {self.docs_directory}")
        
        documents = []
        
        # Load .docx files manually since DirectoryLoader doesn't handle them by default
        for docx_file in self.docs_directory.glob("*.docx"):
            try:
                from docx import Document as DocxDocument
                
                doc = DocxDocument(docx_file)
                content_parts = []
                
                for paragraph in doc.paragraphs:
                    if paragraph.text.strip():
                        content_parts.append(paragraph.text.strip())
                
                content = "\n".join(content_parts)
                
                if content.strip():  # Only add non-empty files
                    doc_obj = Document(
                        page_content=content,
                        metadata={"source": str(docx_file), "filename": docx_file.name}
                    )
                    documents.append(doc_obj)
                    logger.info(f"✅ Loaded {docx_file.name}: {len(content)} characters")
                
            except Exception as e:
                logger.error(f"❌ Error loading {docx_file.name}: {e}")
        
        if not documents:
            raise ValueError(f"No valid documents found in {self.docs_directory}")
        
        logger.info(f"✅ Successfully loaded {len(documents)} documents")
        return documents
    
    def split_documents(self, documents: List[Document]) -> List[Document]:
        """
        Split documents into chunks using RecursiveCharacterTextSplitter
        
        Args:
            documents: List of Document objects to split
            
        Returns:
            List of split Document chunks
        """
        logger.info("Splitting documents into chunks...")
        
        if not self.text_splitter:
            self.setup_text_splitter()
        
        # Split documents
        split_docs = self.text_splitter.split_documents(documents)
        
        logger.info(f"✅ Split into {len(split_docs)} chunks")
        
        # Log chunk statistics
        chunk_sizes = [len(chunk.page_content) for chunk in split_docs]
        avg_chunk_size = sum(chunk_sizes) / len(chunk_sizes)
        min_chunk_size = min(chunk_sizes)
        max_chunk_size = max(chunk_sizes)
        
        logger.info(f"  📊 Chunk statistics:")
        logger.info(f"    Average size: {avg_chunk_size:.0f} characters")
        logger.info(f"    Min size: {min_chunk_size} characters")
        logger.info(f"    Max size: {max_chunk_size} characters")
        
        return split_docs
    
    def create_vectorstore(self, documents: List[Document]) -> Chroma:
        """
        Create ChromaDB vector store and store document chunks with embeddings
        
        Args:
            documents: List of Document chunks to store
            
        Returns:
            ChromaDB vector store instance
        """
        logger.info("Creating ChromaDB vector store...")
        
        if not self.embeddings:
            self.setup_embeddings()
        
        # Create persistent directory
        self.persist_directory.mkdir(exist_ok=True)
        
        # Create ChromaDB vector store
        self.vectorstore = Chroma.from_documents(
            documents=documents,
            embedding=self.embeddings,
            collection_name=self.collection_name,
            persist_directory=str(self.persist_directory)
        )
        
        # Persist the vector store
        self.vectorstore.persist()
        
        logger.info(f"✅ Vector store created and persisted to: {self.persist_directory}")
        logger.info(f"✅ Collection name: {self.collection_name}")
        
        return self.vectorstore
    
    def get_collection_stats(self) -> Dict[str, Any]:
        """
        Get statistics about the created collection
        
        Returns:
            Dictionary with collection statistics
        """
        if not self.vectorstore:
            return {"error": "Vector store not initialized"}
        
        try:
            # Get collection info
            collection = self.vectorstore._collection
            
            stats = {
                "collection_name": self.collection_name,
                "persist_directory": str(self.persist_directory),
                "total_documents": collection.count(),
                "embedding_model": "text-embedding-3-small",
                "chunk_size": self.chunk_size,
                "chunk_overlap": self.chunk_overlap
            }
            
            return stats
            
        except Exception as e:
            logger.error(f"Error getting collection stats: {e}")
            return {"error": str(e)}
    
    def run_pipeline(self) -> Dict[str, Any]:
        """
        Run the complete data loading and embedding pipeline
        
        Returns:
            Dictionary with pipeline results and statistics
        """
        logger.info("🚀 Starting SOM Data Loading and Embedding Pipeline")
        logger.info("=" * 60)
        
        try:
            # Step 1: Load documents
            logger.info("Step 1: Loading documents...")
            documents = self.load_documents()
            
            # Step 2: Split documents
            logger.info("Step 2: Splitting documents...")
            split_docs = self.split_documents(documents)
            
            # Step 3: Setup embeddings
            logger.info("Step 3: Setting up embeddings...")
            self.setup_embeddings()
            
            # Step 4: Create vector store
            logger.info("Step 4: Creating vector store...")
            vectorstore = self.create_vectorstore(split_docs)
            
            # Get final statistics
            stats = self.get_collection_stats()
            
            logger.info("=" * 60)
            logger.info("🎉 Pipeline completed successfully!")
            logger.info("=" * 60)
            
            return {
                "success": True,
                "documents_loaded": len(documents),
                "chunks_created": len(split_docs),
                "collection_stats": stats
            }
            
        except Exception as e:
            logger.error(f"❌ Pipeline failed: {e}")
            return {
                "success": False,
                "error": str(e)
            }

def main():
    """
    Main function to run the SOM Data Loading and Embedding Pipeline
    """
    print("🧠 MindShift SOM Data Loading and Embedding Pipeline")
    print("=" * 60)
    print("Mirrors the n8n 'Load Data Flow' using LangChain")
    print("=" * 60)
    
    # Check for required environment variables
    if not os.getenv("OPENAI_API_KEY"):
        print("❌ Error: OPENAI_API_KEY environment variable not set")
        print("Please set your OpenAI API key in a .env file")
        print("You can get an API key from: https://platform.openai.com/api-keys")
        return
    
    # Initialize the data loader
    try:
        data_loader = SOMDataLoader(
            docs_directory="./som_documents",
            collection_name="som_mindshift",
            persist_directory="./chroma_db",
            chunk_size=1000,
            chunk_overlap=200
        )
        
        # Run the pipeline
        results = data_loader.run_pipeline()
        
        if results["success"]:
            print("\n✅ SUCCESS: Data loading and embedding completed!")
            print(f"📊 Documents loaded: {results['documents_loaded']}")
            print(f"📊 Chunks created: {results['chunks_created']}")
            print(f"📊 Collection: {results['collection_stats']['collection_name']}")
            print(f"📊 Total documents in collection: {results['collection_stats']['total_documents']}")
            
            print("\n🎯 Next steps:")
            print("1. Use the vector store for RAG queries")
            print("2. Run retrieval and generation pipeline")
            print("3. Test with sample queries")
            
        else:
            print(f"\n❌ FAILED: {results['error']}")
            
    except Exception as e:
        print(f"\n❌ Error: {e}")

if __name__ == "__main__":
    main()
