import os
import logging
from typing import List, Dict, Any
from pathlib import Path
import chromadb
from chromadb.config import Settings

from llama_index.core import (
    VectorStoreIndex, 
    SimpleDirectoryReader, 
    StorageContext,
    Settings as LlamaIndexSettings
)
from llama_index.vector_stores.chroma import ChromaVectorStore
from llama_index.embeddings.openai import OpenAIEmbedding
from llama_index.llms.openai import OpenAI
from llama_index.core.node_parser import SentenceSplitter
from llama_index.core.schema import Document
from llama_index.core.query_engine import RetrieverQueryEngine
from llama_index.core.retrievers import VectorIndexRetriever
from llama_index.core.response_synthesizers import ResponseMode

import streamlit as st
from dotenv import load_dotenv

# Load environment variables
load_dotenv()

# Configure logging
logging.basicConfig(level=logging.INFO)
logger = logging.getLogger(__name__)

class MindShiftRAG:
    """
    MindShift RAG system for NLP coaching using SOM patterns
    """
    
    def __init__(self, persist_dir: str = "./chroma_db"):
        self.persist_dir = persist_dir
        self.chroma_client = None
        self.chroma_collection = None
        self.vector_store = None
        self.index = None
        self.query_engine = None
        
        # Initialize LlamaIndex settings
        LlamaIndexSettings.embed_model = OpenAIEmbedding(
            model="text-embedding-3-small",
            api_key=os.getenv("OPENAI_API_KEY")
        )
        LlamaIndexSettings.llm = OpenAI(
            model="gpt-4",
            api_key=os.getenv("OPENAI_API_KEY"),
            temperature=0.7
        )
        
    def setup_chromadb(self):
        """Initialize ChromaDB client and collection"""
        try:
            self.chroma_client = chromadb.PersistentClient(
                path=self.persist_dir,
                settings=Settings(
                    anonymized_telemetry=False,
                    allow_reset=True
                )
            )
            
            # Create or get collection
            self.chroma_collection = self.chroma_client.get_or_create_collection(
                name="mindshift_som_patterns",
                metadata={"description": "SOM patterns and conversation examples"}
            )
            
            self.vector_store = ChromaVectorStore(
                chroma_collection=self.chroma_collection
            )
            
            logger.info("ChromaDB setup completed successfully")
            
        except Exception as e:
            logger.error(f"Error setting up ChromaDB: {e}")
            raise
    
    def load_documents(self, docs_dir: str) -> List[Document]:
        """Load and parse .docx and .txt files from the specified directory"""
        try:
            documents = []
            docs_path = Path(docs_dir)
            
            if not docs_path.exists():
                raise FileNotFoundError(f"Documents directory {docs_dir} not found")
            
            # Load .txt files
            for txt_file in docs_path.glob("*.txt"):
                try:
                    with open(txt_file, 'r', encoding='utf-8') as f:
                        content = f.read()
                    
                    if content.strip():  # Only add non-empty files
                        doc = Document(
                            text=content,
                            metadata={"filename": txt_file.name, "file_type": "txt"}
                        )
                        documents.append(doc)
                        logger.info(f"Loaded {txt_file.name}: {len(content)} characters")
                        
                except Exception as e:
                    logger.error(f"Error loading {txt_file.name}: {e}")
            
            # Load .docx files if python-docx is available
            try:
                from docx import Document as DocxDocument
                
                for docx_file in docs_path.glob("*.docx"):
                    try:
                        doc = DocxDocument(docx_file)
                        content_parts = []
                        
                        for paragraph in doc.paragraphs:
                            if paragraph.text.strip():
                                content_parts.append(paragraph.text.strip())
                        
                        content = "\n".join(content_parts)
                        
                        if content.strip():  # Only add non-empty files
                            doc_obj = Document(
                                text=content,
                                metadata={"filename": docx_file.name, "file_type": "docx"}
                            )
                            documents.append(doc_obj)
                            logger.info(f"Loaded {docx_file.name}: {len(content)} characters")
                            
                    except Exception as e:
                        logger.error(f"Error loading {docx_file.name}: {e}")
                        
            except ImportError:
                logger.warning("python-docx not installed. Install with: pip install python-docx")
            
            if not documents:
                raise ValueError(f"No valid documents found in {docs_dir}")
            
            logger.info(f"Successfully loaded {len(documents)} documents")
            return documents
            
        except Exception as e:
            logger.error(f"Error loading documents: {e}")
            raise
    
    def create_index(self, documents: List[Document]):
        """Create vector index from documents"""
        try:
            # Create storage context
            storage_context = StorageContext.from_defaults(
                vector_store=self.vector_store
            )
            
            # Create index
            self.index = VectorStoreIndex.from_documents(
                documents,
                storage_context=storage_context,
                transformations=[
                    SentenceSplitter(
                        chunk_size=512,
                        chunk_overlap=50
                    )
                ]
            )
            
            logger.info("Vector index created successfully")
            
        except Exception as e:
            logger.error(f"Error creating index: {e}")
            raise
    
    def setup_query_engine(self):
        """Setup the query engine for retrieval and generation"""
        try:
            # Create retriever
            retriever = VectorIndexRetriever(
                index=self.index,
                similarity_top_k=5
            )
            
            # Create query engine
            self.query_engine = RetrieverQueryEngine.from_args(
                retriever=retriever,
                response_mode=ResponseMode.COMPACT,
                node_postprocessors=[]
            )
            
            logger.info("Query engine setup completed")
            
        except Exception as e:
            logger.error(f"Error setting up query engine: {e}")
            raise
    
    def initialize_system(self, docs_dir: str):
        """Initialize the complete RAG system"""
        logger.info("Initializing MindShift RAG system...")
        
        # Setup ChromaDB
        self.setup_chromadb()
        
        # Load documents
        documents = self.load_documents(docs_dir)
        
        # Create index
        self.create_index(documents)
        
        # Setup query engine
        self.setup_query_engine()
        
        logger.info("MindShift RAG system initialized successfully!")
    
    def query(self, user_input: str) -> str:
        """Query the RAG system with user input"""
        try:
            # First, retrieve relevant documents
            retriever = VectorIndexRetriever(
                index=self.index,
                similarity_top_k=5
            )
            
            # Get relevant nodes
            nodes = retriever.retrieve(user_input)
            
            # Extract the retrieved content
            retrieved_content = "\n\n".join([node.text for node in nodes])
            
            logger.info(f"Retrieved {len(nodes)} relevant documents")
            logger.info(f"Retrieved content length: {len(retrieved_content)} characters")
            
            # Create a more specific prompt for SOM pattern coaching
            coaching_prompt = f"""
            You are MindShift, an expert NLP coach specializing in Sleight of Mouth (SOM) patterns. 
            The user has shared a limiting belief: "{user_input}"
            
            Here are the relevant SOM patterns and conversation examples from the knowledge base:
            
            {retrieved_content}
            
            Based on this retrieved information, please:
            
            1. Identify 2-3 relevant SOM patterns from the knowledge base that could challenge this belief
            2. Formulate 2-3 powerful questions using these specific patterns
            3. Maintain a supportive, conversational tone like in the examples
            4. Reference specific patterns by name when appropriate
            5. Use the conversation style and approach shown in the examples
            
            Focus on helping the user reframe their limiting belief into a more empowering perspective using the exact patterns and techniques from the knowledge base.
            """
            
            # Use the LLM directly with the retrieved content
            response = LlamaIndexSettings.llm.complete(coaching_prompt)
            return str(response)
            
        except Exception as e:
            logger.error(f"Error querying system: {e}")
            return f"I apologize, but I encountered an error while processing your request: {e}"
    
    def get_collection_stats(self) -> Dict[str, Any]:
        """Get statistics about the collection"""
        try:
            count = self.chroma_collection.count()
            return {
                "total_documents": count,
                "collection_name": self.chroma_collection.name,
                "persist_directory": self.persist_dir
            }
        except Exception as e:
            logger.error(f"Error getting collection stats: {e}")
            return {"error": str(e)}

def main():
    """Main function to run the MindShift RAG system"""
    
    # Check for OpenAI API key
    if not os.getenv("OPENAI_API_KEY"):
        print("Error: OPENAI_API_KEY environment variable not set")
        print("Please set your OpenAI API key in a .env file or environment variable")
        return
    
    # Initialize the RAG system
    rag_system = MindShiftRAG()
    
    # Set the documents directory (you'll need to download the files here)
    docs_directory = "./som_documents"
    
    # Check if documents directory exists
    if not os.path.exists(docs_directory):
        print(f"Documents directory '{docs_directory}' not found.")
        print("Please download the SOM pattern documents and place them in this directory.")
        return
    
    try:
        # Initialize the system
        rag_system.initialize_system(docs_directory)
        
        # Get and display collection stats
        stats = rag_system.get_collection_stats()
        print(f"Collection Stats: {stats}")
        
        # Interactive chat loop
        print("\n" + "="*60)
        print("🤖 MindShift NLP Coach - Ready to Help!")
        print("="*60)
        print("Share your limiting belief and I'll help you reframe it using SOM patterns.")
        print("Type 'quit' to exit.\n")
        
        while True:
            user_input = input("You: ").strip()
            
            if user_input.lower() in ['quit', 'exit', 'bye']:
                print("Thank you for using MindShift! Take care! 🌟")
                break
            
            if not user_input:
                continue
            
            print("\nMindShift: ", end="")
            response = rag_system.query(user_input)
            print(response)
            print("\n" + "-"*60 + "\n")
            
    except Exception as e:
        logger.error(f"Error in main: {e}")
        print(f"An error occurred: {e}")

if __name__ == "__main__":
    main()
